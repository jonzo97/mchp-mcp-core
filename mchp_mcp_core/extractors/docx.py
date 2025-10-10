"""
DOCX (Microsoft Word) document extraction.

Extracts text, tables, and structure from Word documents using python-docx.
Compatible with ExtractedChunk interface for downstream processing.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any, Optional

from mchp_mcp_core.models import ExtractedChunk
from mchp_mcp_core.utils.logger import get_logger

logger = get_logger(__name__)


class DOCXExtractor:
    """Extract content from Microsoft Word documents."""

    def __init__(self, config: Optional[Any] = None):
        """Initialize DOCX extractor."""
        self.config = config or {}

    def extract_document(
        self,
        file_path: Path | str,
        doc_id: str,
        title: Optional[str] = None
    ) -> list[ExtractedChunk]:
        """
        Extract chunks from DOCX file.

        Args:
            file_path: Path to DOCX file
            doc_id: Unique document identifier
            title: Document title (extracted from properties if not provided)

        Returns:
            List of ExtractedChunk objects

        Example:
            >>> extractor = DOCXExtractor()
            >>> chunks = extractor.extract_document("doc.docx", "doc_001")
            >>> print(f"Extracted {len(chunks)} chunks")
        """
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return []

        file_path = Path(file_path) if isinstance(file_path, str) else file_path

        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return []

        try:
            doc = Document(file_path)
            chunks: list[ExtractedChunk] = []

            # Extract title from properties if not provided
            if title is None:
                title = doc.core_properties.title or file_path.stem

            # Extract paragraphs
            current_section = "Introduction"
            current_text = []
            chunk_id = 0

            for para in doc.paragraphs:
                text = para.text.strip()

                if not text:
                    continue

                # Detect section headers (style-based)
                if para.style.name.startswith('Heading'):
                    # Save previous section
                    if current_text:
                        chunk = self._create_chunk(
                            doc_id=doc_id,
                            title=title,
                            source_path=str(file_path),
                            section=current_section,
                            chunk_id=chunk_id,
                            text="\n\n".join(current_text),
                            chunk_type="text"
                        )
                        chunks.append(chunk)
                        chunk_id += 1
                        current_text = []

                    # Start new section
                    current_section = text
                else:
                    current_text.append(text)

            # Save final section
            if current_text:
                chunk = self._create_chunk(
                    doc_id=doc_id,
                    title=title,
                    source_path=str(file_path),
                    section=current_section,
                    chunk_id=chunk_id,
                    text="\n\n".join(current_text),
                    chunk_type="text"
                )
                chunks.append(chunk)

            # Extract tables
            for i, table in enumerate(doc.tables):
                table_text = self._extract_table_text(table)
                if table_text:
                    chunk = self._create_chunk(
                        doc_id=doc_id,
                        title=title,
                        source_path=str(file_path),
                        section="Tables",
                        chunk_id=len(chunks),
                        text=table_text,
                        chunk_type="table"
                    )
                    chunks.append(chunk)

            logger.info(f"Extracted {len(chunks)} chunks from {file_path.name}")
            return chunks

        except Exception as e:
            logger.error(f"Error extracting {file_path}: {e}")
            return []

    def _extract_table_text(self, table) -> str:
        """Extract text from table in markdown format."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))

        if rows:
            # Add separator after header
            if len(rows) > 1:
                separator = " | ".join(["---"] * len(rows[0].split(" | ")))
                return rows[0] + "\n" + separator + "\n" + "\n".join(rows[1:])
            return "\n".join(rows)

        return ""

    def _create_chunk(
        self,
        doc_id: str,
        title: str,
        source_path: str,
        section: str,
        chunk_id: int,
        text: str,
        chunk_type: str
    ) -> ExtractedChunk:
        """Create ExtractedChunk from extracted content."""
        return ExtractedChunk(
            doc_id=doc_id,
            title=title,
            source=source_path,
            page_or_slide=0,  # DOCX doesn't have pages in the same way
            chunk_id=chunk_id,
            section_hierarchy=section,
            content=text,
            chunk_type=chunk_type,
            confidence=1.0
        )
